import os
import uuid
from pathlib import Path
from typing import Dict, Any, Optional
# from docling.document_converter import DocumentConverter
from config import Config


class DoclingService:
    def __init__(self):
        # self.converter = DocumentConverter()
        Config.create_directories()

    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from uploaded document using PyPDF (replacing Docling)

        Args:
            file_path: Path to the uploaded file

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            text_content = ""
            html_content = ""
            metadata = {
                "file_type": "document",
                "file_extension": file_ext,
                "file_size": os.path.getsize(file_path),
                "extraction_method": "pypdf_fallback"
            }

            if file_ext == ".txt":
                # Handle plain text files
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
                metadata["file_type"] = "text"
                metadata["extraction_method"] = "direct_read"

            elif file_ext == ".pdf":
                # Handle PDF files using pypdf
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    text_pages = []
                    for page in reader.pages:
                        text_pages.append(page.extract_text())
                    text_content = "\n\n".join(text_pages)
                    metadata["page_count"] = len(reader.pages)
                    metadata["extraction_method"] = "pypdf"
                    
                    # Simple HTML generation for visualization
                    html_content = f"<div class='document-content'>"
                    for i, page_text in enumerate(text_pages):
                        html_content += f"<div class='page' id='page-{i+1}'><h3>Page {i+1}</h3><p>{page_text.replace(chr(10), '<br>')}</p></div><hr>"
                    html_content += "</div>"
                    
                except ImportError:
                    raise ImportError("pypdf is not installed. Please install it with: pip install pypdf")
                except Exception as e:
                    raise Exception(f"PDF processing failed: {str(e)}")

            elif file_ext in [".docx", ".doc"]:
                # Handle DOCX files using python-docx
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text_content = "\n".join([para.text for para in doc.paragraphs])
                    metadata["extraction_method"] = "python-docx"
                    
                    # Simple HTML generation
                    html_content = "<div class='document-content'>"
                    for para in doc.paragraphs:
                        if para.text.strip():
                            html_content += f"<p>{para.text}</p>"
                    html_content += "</div>"
                except ImportError:
                     raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
                except Exception as e:
                    raise Exception(f"DOCX processing failed: {str(e)}")

            else:
                # Fallback for other types (mock/placeholder)
                text_content = f"Content extraction for {file_ext} not supported in this lightweight mode."
                metadata["extraction_method"] = "unsupported"

            return {
                "success": True,
                "text": text_content,
                "html": html_content,
                "metadata": metadata
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "metadata": {"file_path": file_path}
            }

    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """
        Save uploaded file with a unique name and return the path.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Path to saved file
        """
        try:
            # Create uploads directory if it doesn't exist
            Config.UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

            # Generate unique filename
            file_ext = Path(filename).suffix
            unique_filename = f"{uuid.uuid4()}{file_ext}"
            file_path = Config.UPLOADS_DIR / unique_filename

            # Save file with error handling
            try:
                with open(file_path, "wb") as f:
                    f.write(file_content)

                # Verify file was written
                if not file_path.exists():
                    raise IOError(f"Failed to save file: {file_path}")

                return str(file_path)

            except IOError as e:
                raise Exception(f"Error writing file {filename}: {str(e)}")

        except Exception as e:
            raise Exception(f"Error saving uploaded file: {str(e)}")

    def cleanup_file(self, file_path: str) -> bool:
        """
        Clean up uploaded file

        Args:
            file_path: Path to file to remove

        Returns:
            True if file was removed, False otherwise
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
        except Exception:
            pass
        return False

    def validate_file(self, filename: str, file_size: int) -> Dict[str, Any]:
        """
        Validate uploaded file

        Args:
            filename: Original filename
            file_size: File size in bytes

        Returns:
            Validation result
        """
        file_ext = Path(filename).suffix.lower()

        # Check file extension
        if file_ext not in Config.ALLOWED_EXTENSIONS:
            return {
                "valid": False,
                "error": f"File type {file_ext} not allowed. Allowed types: {Config.ALLOWED_EXTENSIONS}"
            }

        # Check file size
        if file_size > Config.MAX_FILE_SIZE:
            return {
                "valid": False,
                "error": f"File size {file_size} exceeds maximum allowed size {Config.MAX_FILE_SIZE}"
            }

        return {"valid": True}
